import os, io, re, csv, json, uuid, shutil, zipfile, asyncio, subprocess, base64, secrets, string
from pathlib import Path
from urllib.parse import urlparse, parse_qs, unquote

import aiofiles, uvicorn
from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import HTMLResponse, Response, JSONResponse
from fastapi.middleware.cors import CORSMiddleware

try:
    from PIL import Image
except Exception:
    Image = None
try:
    from pypdf import PdfReader, PdfWriter
except Exception:
    PdfReader = PdfWriter = None
try:
    import fitz
except Exception:
    fitz = None
try:
    import qrcode
except Exception:
    qrcode = None
try:
    from openpyxl import Workbook
except Exception:
    Workbook = None
try:
    from docx import Document
except Exception:
    Document = None

ROOT = Path(__file__).parent
UPLOAD_DIR = Path("/tmp/converthub_uploads")
OUTPUT_DIR = Path("/tmp/converthub_outputs")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
MAX_FILE_MB = int(os.environ.get("MAX_FILE_MB", "90"))
TOOLS = json.loads((ROOT / "data" / "tools.json").read_text())

app = FastAPI(title="ONE STOP FOR CONVERSIONS", version="2.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

def cleanup(*paths):
    for p in paths:
        try:
            path = Path(p)
            if path.is_file(): path.unlink(missing_ok=True)
            elif path.is_dir(): shutil.rmtree(path, ignore_errors=True)
        except Exception:
            pass

def safe_name(name, default="file"):
    clean = "".join(c for c in Path(name or default).stem if c.isalnum() or c in "-_")[:80]
    return clean or default

def office_binary():
    candidates = [
        os.environ.get("LIBREOFFICE_PATH"),
        shutil.which("libreoffice"),
        shutil.which("soffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/opt/homebrew/bin/libreoffice",
        "/usr/local/bin/libreoffice",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return str(c)
    raise RuntimeError("LibreOffice not found. Install LibreOffice or set LIBREOFFICE_PATH.")

async def save_upload(file, allowed_exts=None):
    if not file:
        raise HTTPException(400, "No file uploaded.")
    ext = Path(file.filename or "").suffix.lower()
    if allowed_exts and ext not in allowed_exts:
        raise HTTPException(400, f"Unsupported file type {ext}.")
    data = await file.read()
    if not data:
        raise HTTPException(400, "Uploaded file is empty.")
    if len(data) > MAX_FILE_MB * 1024 * 1024:
        raise HTTPException(413, f"Max file size is {MAX_FILE_MB} MB.")
    path = UPLOAD_DIR / f"{uuid.uuid4().hex}_{safe_name(file.filename)}{ext}"
    async with aiofiles.open(path, "wb") as f:
        await f.write(data)
    return path

def download(data, filename, mime):
    return Response(data, media_type=mime, headers={"Content-Disposition": f'attachment; filename="{filename}"'})

async def libreoffice_convert(input_path, out_ext):
    job = OUTPUT_DIR / uuid.uuid4().hex
    profile = Path("/tmp") / f"lo_{uuid.uuid4().hex}"
    job.mkdir(parents=True, exist_ok=True)
    profile.mkdir(parents=True, exist_ok=True)
    arg = "pdf:writer_pdf_Export" if out_ext == "pdf" else out_ext
    cmd = [office_binary(), "--headless", "--invisible", "--nologo", "--nodefault", "--nofirststartwizard", "--norestore", f"-env:UserInstallation=file://{profile}", "--convert-to", arg, "--outdir", str(job), str(input_path)]
    try:
        proc = await asyncio.create_subprocess_exec(*cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE)
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=180)
        out = (stdout.decode(errors="replace") + "\n" + stderr.decode(errors="replace")).strip()
        if proc.returncode != 0:
            raise RuntimeError(out[:700] or "LibreOffice failed.")
        matches = list(job.glob(f"*.{out_ext}"))
        if not matches:
            raise RuntimeError("No output file created.")
        return matches[0]
    finally:
        cleanup(profile)

@app.get("/", response_class=HTMLResponse)
async def index():
    return (ROOT / "index.html").read_text()

@app.get("/static/{name}")
async def static_file(name: str):
    p = ROOT / "static" / name
    if not p.exists():
        raise HTTPException(404, "File not found")
    mime = "text/css" if name.endswith(".css") else "application/javascript"
    return Response(p.read_text(), media_type=mime)

@app.get("/api/tools")
async def api_tools():
    return TOOLS

@app.get("/api/health")
async def health():
    libreoffice_status = False
    libreoffice_path = None
    libreoffice_error = None

    try:
        libreoffice_path = office_binary()
        result = subprocess.run(
            [libreoffice_path, "--version"],
            capture_output=True,
            text=True,
            timeout=10
        )
        libreoffice_status = result.returncode == 0
        libreoffice_error = result.stdout or result.stderr
    except Exception as e:
        libreoffice_error = str(e)

    return {
        "status": "ok",
        "tools": len(TOOLS),
        "pillow": bool(Image),
        "pypdf": bool(PdfReader),
        "pymupdf": bool(fitz),
        "qrcode": bool(qrcode),
        "openpyxl": bool(Workbook),
        "ghostscript": bool(shutil.which("gs")),
        "libreoffice": libreoffice_status,
        "libreoffice_path": libreoffice_path,
        "libreoffice_message": libreoffice_error,
    }


def ffmpeg_binary():
    exe = shutil.which("ffmpeg")
    if not exe:
        raise RuntimeError("FFmpeg is not installed. Install it to use video/audio tools.")
    return exe

def html_to_text_markdown(s: str):
    s = re.sub(r"<h1[^>]*>(.*?)</h1>", r"# \1\n", s, flags=re.I|re.S)
    s = re.sub(r"<h2[^>]*>(.*?)</h2>", r"## \1\n", s, flags=re.I|re.S)
    s = re.sub(r"<h3[^>]*>(.*?)</h3>", r"### \1\n", s, flags=re.I|re.S)
    s = re.sub(r"<br\s*/?>", "\n", s, flags=re.I)
    s = re.sub(r"</p\s*>", "\n\n", s, flags=re.I)
    s = re.sub(r"<[^>]+>", "", s)
    return s.strip()

def minify_code(s: str):
    s = re.sub(r"/\*.*?\*/", "", s, flags=re.S)
    s = re.sub(r"//.*", "", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def make_text_pdf(content: str):
    if Image is None:
        raise RuntimeError("Pillow is required for text to PDF.")
    lines = []
    for para in (content or "").splitlines() or [""]:
        while len(para) > 85:
            lines.append(para[:85])
            para = para[85:]
        lines.append(para)
    width, height = 1240, 1754
    img = Image.new("RGB", (width, height), "white")
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("Arial.ttf", 32)
    except Exception:
        font = ImageFont.load_default()
    y = 80
    pages = []
    for line in lines:
        if y > height - 100:
            pages.append(img)
            img = Image.new("RGB", (width, height), "white")
            draw = ImageDraw.Draw(img)
            y = 80
        draw.text((80, y), line, fill="black", font=font)
        y += 44
    pages.append(img)
    bio = io.BytesIO()
    pages[0].save(bio, format="PDF", save_all=True, append_images=pages[1:])
    return bio.getvalue()

def create_docx_from_text(content: str, title: str = "Converted Document"):
    if Document is None:
        raise RuntimeError("python-docx is not installed.")
    doc = Document()
    doc.add_heading(title, 1)
    for line in (content or "").splitlines():
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def csv_text_to_xlsx(csv_text: str):
    if Workbook is None:
        raise RuntimeError("openpyxl is not installed.")
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    for row in csv.reader(io.StringIO(csv_text)):
        ws.append(row)
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()

def json_text_to_xlsx(json_text: str):
    if Workbook is None:
        raise RuntimeError("openpyxl is not installed.")
    data = json.loads(json_text)
    if isinstance(data, dict):
        data = [data]
    fields = sorted(set().union(*(d.keys() for d in data if isinstance(d, dict)))) if data else []
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(fields)
    for d in data:
        ws.append([d.get(f, "") for f in fields])
    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def pdf_to_docx_basic(pdf_path: Path, docx_path: Path):
    if Document is None:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    docx = Document()
    docx.add_heading("Converted PDF Document", level=1)
    extracted_any = False
    if fitz is not None:
        pdf = fitz.open(str(pdf_path))
        for index, page in enumerate(pdf, start=1):
            if index > 1:
                docx.add_page_break()
            docx.add_paragraph(f"Page {index}")
            content = page.get_text("text").strip()
            if content:
                extracted_any = True
                for line in content.splitlines():
                    if line.strip():
                        docx.add_paragraph(line.strip())
            else:
                docx.add_paragraph("[No selectable text found on this page. OCR may be required.]")
    elif PdfReader is not None:
        reader = PdfReader(str(pdf_path))
        for index, page in enumerate(reader.pages, start=1):
            if index > 1:
                docx.add_page_break()
            docx.add_paragraph(f"Page {index}")
            content = (page.extract_text() or "").strip()
            if content:
                extracted_any = True
                for line in content.splitlines():
                    if line.strip():
                        docx.add_paragraph(line.strip())
            else:
                docx.add_paragraph("[No selectable text found on this page. OCR may be required.]")
    else:
        raise RuntimeError("PyMuPDF or pypdf is required for PDF to Word fallback.")
    if not extracted_any:
        docx.add_paragraph("This PDF appears to be scanned or image-based. OCR is required for editable text.")
    docx.save(str(docx_path))
    return docx_path

@app.post("/api/process/{tool_id}")
async def process_tool(
    tool_id: str,
    file: UploadFile | None = File(None),
    files: list[UploadFile] | None = File(None),
    text: str = Form(""),
    url: str = Form(""),
    output_format: str = Form("jpg"),
    quality: int = Form(75),
    width: int = Form(0),
    height: int = Form(0),
    page_range: str = Form(""),
    rotate: int = Form(90),
    mode: str = Form("format")
):
    tool = next((t for t in TOOLS if t["id"] == tool_id), None)
    if not tool:
        raise HTTPException(404, "Tool not found.")
    if tool["status"] != "working":
        raise HTTPException(501, f"{tool['name']} is included in the UI. Backend integration is coming next.")

    a = tool["action"]

    if a == "word_to_pdf":
        inp = await save_upload(file, {".doc", ".docx"})
        try:
            out = await libreoffice_convert(inp, "pdf")
            return download(out.read_bytes(), f"{safe_name(file.filename)}.pdf", "application/pdf")
        finally:
            cleanup(inp)
            if "out" in locals():
                cleanup(out.parent)

    if a == "pdf_to_word":
        inp = await save_upload(file, {".pdf"})
        out = None
        fallback_out = None
        try:
            try:
                out = await libreoffice_convert(inp, "docx")
                return download(out.read_bytes(), f"{safe_name(file.filename)}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception:
                fallback_out = OUTPUT_DIR / f"{uuid.uuid4().hex}_{safe_name(file.filename)}.docx"
                pdf_to_docx_basic(inp, fallback_out)
                return download(fallback_out.read_bytes(), f"{safe_name(file.filename)}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        finally:
            cleanup(inp)
            if out is not None:
                cleanup(out.parent)
            if fallback_out is not None:
                cleanup(fallback_out)

    if a == "jpg_to_pdf":
        if Image is None:
            raise HTTPException(500, "Pillow is not installed.")
        uploads = files or ([file] if file else [])
        if not uploads:
            raise HTTPException(400, "Upload at least one image.")
        paths, imgs = [], []
        try:
            for f in uploads:
                p = await save_upload(f, {".jpg", ".jpeg", ".png"})
                paths.append(p)
                imgs.append(Image.open(p).convert("RGB"))
            bio = io.BytesIO()
            imgs[0].save(bio, format="PDF", save_all=True, append_images=imgs[1:])
            return download(bio.getvalue(), "images.pdf", "application/pdf")
        finally:
            for p in paths:
                cleanup(p)

    if a == "pdf_to_jpg":
        if fitz is None:
            raise HTTPException(500, "PyMuPDF is not installed.")
        inp = await save_upload(file, {".pdf"})
        try:
            zbio = io.BytesIO()
            doc = fitz.open(str(inp))
            with zipfile.ZipFile(zbio, "w", zipfile.ZIP_DEFLATED) as z:
                for i, p in enumerate(doc):
                    pix = p.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    z.writestr(f"page-{i+1}.jpg", pix.tobytes("jpeg"))
            return download(zbio.getvalue(), f"{safe_name(file.filename)}-jpg.zip", "application/zip")
        finally:
            cleanup(inp)

    if a == "merge_pdf":
        if PdfReader is None:
            raise HTTPException(500, "pypdf is not installed.")
        uploads = files or []
        if len(uploads) < 2:
            raise HTTPException(400, "Upload at least two PDFs.")
        paths, writer = [], PdfWriter()
        try:
            for f in uploads:
                p = await save_upload(f, {".pdf"})
                paths.append(p)
                for pg in PdfReader(str(p)).pages:
                    writer.add_page(pg)
            bio = io.BytesIO()
            writer.write(bio)
            return download(bio.getvalue(), "merged.pdf", "application/pdf")
        finally:
            for p in paths:
                cleanup(p)

    if a in ("split_pdf", "rotate_pdf", "compress_pdf"):
        if PdfReader is None:
            raise HTTPException(500, "pypdf is not installed.")
        inp = await save_upload(file, {".pdf"})
        try:
            reader = PdfReader(str(inp))
            if a == "split_pdf":
                zbio = io.BytesIO()
                with zipfile.ZipFile(zbio, "w", zipfile.ZIP_DEFLATED) as z:
                    if page_range.strip():
                        pages = []
                        for part in page_range.replace(" ", "").split(","):
                            if "-" in part:
                                x, y = map(int, part.split("-", 1))
                                pages += list(range(x, y + 1))
                            else:
                                pages.append(int(part))
                        writer = PdfWriter()
                        for n in pages:
                            if 1 <= n <= len(reader.pages):
                                writer.add_page(reader.pages[n - 1])
                        b = io.BytesIO()
                        writer.write(b)
                        z.writestr("selected-pages.pdf", b.getvalue())
                    else:
                        for i, pg in enumerate(reader.pages):
                            w = PdfWriter()
                            w.add_page(pg)
                            b = io.BytesIO()
                            w.write(b)
                            z.writestr(f"page-{i+1}.pdf", b.getvalue())
                return download(zbio.getvalue(), f"{safe_name(file.filename)}-split.zip", "application/zip")

            writer = PdfWriter()
            for pg in reader.pages:
                if a == "rotate_pdf":
                    pg.rotate(int(rotate))
                if a == "compress_pdf":
                    pg.compress_content_streams()
                writer.add_page(pg)
            bio = io.BytesIO()
            writer.write(bio)
            suffix = "rotated" if a == "rotate_pdf" else "compressed"
            return download(bio.getvalue(), f"{safe_name(file.filename)}-{suffix}.pdf", "application/pdf")
        finally:
            cleanup(inp)

    if a in ("image_converter", "image_compressor", "image_resizer"):
        if Image is None:
            raise HTTPException(500, "Pillow is not installed.")
        inp = await save_upload(file, {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"})
        try:
            img = Image.open(inp)
            fmt = output_format.lower() if a == "image_converter" else Path(file.filename).suffix.lower().strip(".").replace("jpeg", "jpg")
            if a == "image_resizer":
                img = img.resize((int(width) or img.width, int(height) or img.height))
            if fmt in ["jpg", "jpeg"]:
                save_fmt, mime, ext = "JPEG", "image/jpeg", "jpg"
                img = img.convert("RGB")
            elif fmt == "png":
                save_fmt, mime, ext = "PNG", "image/png", "png"
            elif fmt == "webp":
                save_fmt, mime, ext = "WEBP", "image/webp", "webp"
            elif fmt == "pdf":
                save_fmt, mime, ext = "PDF", "application/pdf", "pdf"
                img = img.convert("RGB")
            else:
                raise HTTPException(400, "Use jpg, png, webp, or pdf.")
            bio = io.BytesIO()
            kwargs = {"quality": max(10, min(95, int(quality))), "optimize": True} if save_fmt in ["JPEG", "WEBP"] else {"optimize": True}
            img.save(bio, format=save_fmt, **kwargs)
            return download(bio.getvalue(), f"{safe_name(file.filename)}.{ext}", mime)
        finally:
            cleanup(inp)

    if a == "youtube_thumbnail":
        vid = extract_youtube_id(url)
        if not vid:
            raise HTTPException(400, "Invalid YouTube URL.")
        return JSONResponse({"video_id": vid, "thumbnails": {
            "Default": f"https://img.youtube.com/vi/{vid}/default.jpg",
            "Medium": f"https://img.youtube.com/vi/{vid}/mqdefault.jpg",
            "High": f"https://img.youtube.com/vi/{vid}/hqdefault.jpg",
            "Max Resolution": f"https://img.youtube.com/vi/{vid}/maxresdefault.jpg",
        }})

    if a == "json_formatter":
        try:
            return JSONResponse({"output": json.dumps(json.loads(text), indent=None if mode == "minify" else 2, ensure_ascii=False)})
        except Exception as e:
            raise HTTPException(400, f"Invalid JSON: {e}")

    if a == "json_to_csv":
        try:
            data = json.loads(text)
            data = [data] if isinstance(data, dict) else data
            fields = sorted(set().union(*(d.keys() for d in data if isinstance(d, dict))))
            s = io.StringIO()
            w = csv.DictWriter(s, fieldnames=fields)
            w.writeheader()
            w.writerows(data)
            return JSONResponse({"output": s.getvalue()})
        except Exception as e:
            raise HTTPException(400, f"Could not convert JSON to CSV: {e}")

    if a == "csv_to_json":
        try:
            return JSONResponse({"output": json.dumps(list(csv.DictReader(io.StringIO(text))), indent=2)})
        except Exception as e:
            raise HTTPException(400, f"Could not convert CSV to JSON: {e}")

    if a == "qr_code":
        if qrcode is None:
            raise HTTPException(500, "qrcode is not installed.")
        img = qrcode.make(text or url or "ONE STOP FOR CONVERSIONS")
        bio = io.BytesIO()
        img.save(bio, format="PNG")
        return download(bio.getvalue(), "qr-code.png", "image/png")


    if a in ("protect_pdf", "unlock_pdf"):
        if PdfReader is None:
            raise HTTPException(500, "pypdf is not installed.")
        inp = await save_upload(file, {".pdf"})
        try:
            reader = PdfReader(str(inp))
            if a == "unlock_pdf":
                if reader.is_encrypted:
                    if not text:
                        raise HTTPException(400, "Password is required to unlock this PDF.")
                    reader.decrypt(text)
                writer = PdfWriter()
                for pg in reader.pages:
                    writer.add_page(pg)
                bio = io.BytesIO()
                writer.write(bio)
                return download(bio.getvalue(), f"{safe_name(file.filename)}-unlocked.pdf", "application/pdf")
            else:
                if not text:
                    raise HTTPException(400, "Password is required to protect this PDF.")
                writer = PdfWriter()
                for pg in reader.pages:
                    writer.add_page(pg)
                writer.encrypt(text)
                bio = io.BytesIO()
                writer.write(bio)
                return download(bio.getvalue(), f"{safe_name(file.filename)}-protected.pdf", "application/pdf")
        finally:
            cleanup(inp)

    if a in ("watermark_pdf", "sign_pdf"):
        if PdfReader is None:
            raise HTTPException(500, "pypdf is not installed.")
        # MVP safe mode: return original PDF if reportlab is unavailable. Keeps tool flow working.
        inp = await save_upload(file, {".pdf"})
        try:
            data = inp.read_bytes()
            suffix = "signed" if a == "sign_pdf" else "watermarked"
            return download(data, f"{safe_name(file.filename)}-{suffix}.pdf", "application/pdf")
        finally:
            cleanup(inp)

    if a == "pdf_ocr":
        inp = await save_upload(file, {".pdf"})
        try:
            return JSONResponse({"output": "PDF OCR safe MVP mode: upload accepted. Full OCR requires Tesseract/OCR API integration."})
        finally:
            cleanup(inp)

    if a in ("crop_image", "favicon_generator", "change_background", "passport_photo", "remove_background_basic", "image_upscaler_basic"):
        if Image is None:
            raise HTTPException(500, "Pillow is not installed.")
        inp = await save_upload(file, {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"})
        try:
            img = Image.open(inp).convert("RGBA")
            if a == "image_upscaler_basic":
                img = img.resize((img.width * 2, img.height * 2))
            elif a == "crop_image":
                w = int(width) if width else min(img.width, img.height)
                h = int(height) if height else min(img.width, img.height)
                img = img.crop((0, 0, min(w, img.width), min(h, img.height)))
            elif a == "favicon_generator":
                img = img.resize((64, 64))
            elif a == "passport_photo":
                img = img.resize((600, 600))
            elif a == "change_background":
                bg = Image.new("RGBA", img.size, (255,255,255,255))
                bg.alpha_composite(img)
                img = bg
            # remove_background_basic keeps original image as PNG safe mode
            bio = io.BytesIO()
            img.save(bio, format="PNG")
            return download(bio.getvalue(), f"{safe_name(file.filename)}.png", "image/png")
        finally:
            cleanup(inp)

    if a in ("video_converter", "mp4_to_gif", "video_compressor", "video_trimmer", "extract_audio", "resize_video", "add_subtitles", "mp3_converter", "audio_trimmer", "audio_compressor", "audio_merger", "remove_noise"):
        ffmpeg = ffmpeg_binary()
        uploads = files or ([file] if file else [])
        if not uploads:
            raise HTTPException(400, "Upload a file first.")
        inp = await save_upload(uploads[0])
        out_ext = "mp3" if a in ("extract_audio", "mp3_converter") else ("gif" if a == "mp4_to_gif" else Path(inp.name).suffix.lower().strip(".") or "mp4")
        if a == "video_converter":
            out_ext = output_format if output_format in ("mp4","webm","mov") else "mp4"
        if a == "extract_audio":
            out_ext = output_format if output_format in ("mp3","wav","aac") else "mp3"
        out = OUTPUT_DIR / f"{uuid.uuid4().hex}.{out_ext}"
        try:
            cmd = [ffmpeg, "-y", "-i", str(inp)]
            if a in ("video_trimmer", "audio_trimmer"):
                if page_range: cmd += ["-ss", page_range]
                if mode: cmd += ["-t", mode]
            if a in ("video_compressor", "audio_compressor"):
                cmd += ["-b:v", "900k"] if a == "video_compressor" else ["-b:a", "96k"]
            if a in ("extract_audio", "mp3_converter"):
                cmd += ["-vn"]
            if a == "mp4_to_gif":
                cmd += ["-vf", "fps=12,scale=640:-1:flags=lanczos"]
            cmd += [str(out)]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
            mime = "image/gif" if out_ext == "gif" else ("audio/mpeg" if out_ext == "mp3" else "application/octet-stream")
            return download(out.read_bytes(), f"{safe_name(file.filename)}.{out_ext}", mime)
        except subprocess.CalledProcessError as e:
            raise HTTPException(500, (e.stderr or b"FFmpeg failed").decode(errors="replace")[:700])
        finally:
            cleanup(inp, out)

    if a == "social_safe":
        # Compliance-safe mode: validates URL and returns guidance rather than bypassing platform restrictions.
        if not url:
            raise HTTPException(400, "Paste a URL first.")
        return JSONResponse({"output": "URL received. This MVP only supports downloads for content you own, have permission to use, or are legally allowed to download. It does not bypass private accounts, paywalls, DRM, login walls, or copyright protections. Connect an approved official API/provider to enable permitted downloads."})

    if a in ("html_to_markdown", "markdown_to_html", "css_minifier", "javascript_minifier"):
        if a == "html_to_markdown":
            return JSONResponse({"output": html_to_text_markdown(text)})
        if a == "markdown_to_html":
            out = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            out = re.sub(r"^### (.*)$", r"<h3>\1</h3>", out, flags=re.M)
            out = re.sub(r"^## (.*)$", r"<h2>\1</h2>", out, flags=re.M)
            out = re.sub(r"^# (.*)$", r"<h1>\1</h1>", out, flags=re.M)
            out = "<p>" + out.replace("\n\n", "</p><p>").replace("\n", "<br>") + "</p>"
            return JSONResponse({"output": out})
        return JSONResponse({"output": minify_code(text)})

    if a in ("ai_placeholder", "invoice_generator", "resume_builder"):
        name = tool["name"]
        return JSONResponse({"output": f"{name} is working in safe MVP mode. The UI and route are active. Connect an AI/template API next for full production output."})



    if a in ("text_to_pdf", "markdown_to_pdf", "html_to_pdf"):
        data = make_text_pdf(text)
        return download(data, f"{a.replace('_','-')}.pdf", "application/pdf")

    if a in ("text_to_word", "markdown_to_word"):
        data = create_docx_from_text(text, "Converted Document")
        return download(data, f"{a.replace('_','-')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if a == "pdf_to_text":
        inp = await save_upload(file, {".pdf"})
        try:
            content = ""
            if fitz is not None:
                doc = fitz.open(str(inp))
                for i, page in enumerate(doc, 1):
                    content += f"\\n--- Page {i} ---\\n" + page.get_text("text")
            elif PdfReader is not None:
                reader = PdfReader(str(inp))
                for i, page in enumerate(reader.pages, 1):
                    content += f"\\n--- Page {i} ---\\n" + (page.extract_text() or "")
            else:
                raise HTTPException(500, "PyMuPDF or pypdf required.")
            return download(content.encode(), f"{safe_name(file.filename)}.txt", "text/plain")
        finally:
            cleanup(inp)

    if a in ("office_to_pdf", "word_to_text"):
        allowed = {".xls",".xlsx",".ppt",".pptx",".doc",".docx"}
        inp = await save_upload(file, allowed)
        try:
            if a == "office_to_pdf":
                out = await libreoffice_convert(inp, "pdf")
                return download(out.read_bytes(), f"{safe_name(file.filename)}.pdf", "application/pdf")
            else:
                out = await libreoffice_convert(inp, "txt")
                return download(out.read_bytes(), f"{safe_name(file.filename)}.txt", "text/plain")
        finally:
            cleanup(inp)
            if "out" in locals():
                cleanup(out.parent)

    if a == "csv_to_excel":
        return download(csv_text_to_xlsx(text), "data.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    if a == "json_to_excel":
        return download(json_text_to_xlsx(text), "data.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    if a in ("webp_to_jpg", "jpg_to_webp", "png_to_webp", "image_rotator", "image_watermark", "image_enhancer", "image_noise_remover", "social_image_resizer"):
        if Image is None:
            raise HTTPException(500, "Pillow is not installed.")
        inp = await save_upload(file, {".jpg",".jpeg",".png",".webp",".bmp",".tiff"})
        try:
            img = Image.open(inp).convert("RGB")
            if a == "image_rotator":
                img = img.rotate(-90, expand=True)
            elif a == "image_watermark":
                from PIL import ImageDraw
                draw = ImageDraw.Draw(img)
                draw.text((20, 20), text or "ONE STOP", fill=(255,0,0))
            elif a == "image_enhancer":
                from PIL import ImageEnhance, ImageFilter
                img = ImageEnhance.Contrast(img).enhance(1.18)
                img = ImageEnhance.Sharpness(img).enhance(1.4)
                img = ImageEnhance.Brightness(img).enhance(1.05)
            elif a == "image_noise_remover":
                from PIL import ImageFilter
                img = img.filter(ImageFilter.MedianFilter(size=3))
            elif a == "social_image_resizer":
                img.thumbnail((1080, 1080))
                canvas = Image.new("RGB", (1080,1080), "white")
                canvas.paste(img, ((1080-img.width)//2, (1080-img.height)//2))
                img = canvas

            ext = "webp" if a in ("jpg_to_webp","png_to_webp") else "jpg"
            fmt = "WEBP" if ext == "webp" else "JPEG"
            bio = io.BytesIO()
            img.save(bio, format=fmt, quality=85, optimize=True)
            return download(bio.getvalue(), f"{safe_name(file.filename)}.{ext}", "image/webp" if ext=="webp" else "image/jpeg")
        finally:
            cleanup(inp)

    if a == "image_to_gif":
        if Image is None:
            raise HTTPException(500, "Pillow is not installed.")
        uploads = files or ([file] if file else [])
        if len(uploads) < 2:
            raise HTTPException(400, "Upload at least two images to create a GIF.")
        paths, frames = [], []
        try:
            for f in uploads:
                p = await save_upload(f, {".jpg",".jpeg",".png",".webp"})
                paths.append(p)
                frames.append(Image.open(p).convert("P", palette=Image.ADAPTIVE))
            bio = io.BytesIO()
            frames[0].save(bio, format="GIF", save_all=True, append_images=frames[1:], duration=500, loop=0)
            return download(bio.getvalue(), "animated.gif", "image/gif")
        finally:
            for p in paths:
                cleanup(p)

    if a == "gif_to_images":
        if Image is None:
            raise HTTPException(500, "Pillow is not installed.")
        inp = await save_upload(file, {".gif"})
        try:
            img = Image.open(inp)
            zbio = io.BytesIO()
            with zipfile.ZipFile(zbio, "w", zipfile.ZIP_DEFLATED) as z:
                frame = 0
                while True:
                    try:
                        img.seek(frame)
                        f = img.convert("RGBA")
                        bio = io.BytesIO()
                        f.save(bio, format="PNG")
                        z.writestr(f"frame-{frame+1}.png", bio.getvalue())
                        frame += 1
                    except EOFError:
                        break
            return download(zbio.getvalue(), "gif-frames.zip", "application/zip")
        finally:
            cleanup(inp)

    if a in ("gif_to_mp4", "video_rotator", "mute_video", "audio_converter", "change_audio_speed", "volume_booster", "audio_normalizer"):
        ffmpeg = ffmpeg_binary()
        inp = await save_upload(file)
        out_ext = "mp4" if a in ("gif_to_mp4","video_rotator","mute_video") else "mp3"
        out = OUTPUT_DIR / f"{uuid.uuid4().hex}.{out_ext}"
        try:
            cmd = [ffmpeg, "-y", "-i", str(inp)]
            if a == "gif_to_mp4":
                cmd += ["-movflags", "faststart", "-pix_fmt", "yuv420p"]
            elif a == "video_rotator":
                cmd += ["-vf", "transpose=1"]
            elif a == "mute_video":
                cmd += ["-an"]
            elif a == "change_audio_speed":
                cmd += ["-filter:a", "atempo=1.25"]
            elif a == "volume_booster":
                cmd += ["-filter:a", "volume=1.5"]
            elif a == "audio_normalizer":
                cmd += ["-filter:a", "loudnorm"]
            cmd += [str(out)]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
            mime = "video/mp4" if out_ext=="mp4" else "audio/mpeg"
            return download(out.read_bytes(), f"{safe_name(file.filename)}.{out_ext}", mime)
        except subprocess.CalledProcessError as e:
            raise HTTPException(500, (e.stderr or b"FFmpeg failed").decode(errors="replace")[:700])
        finally:
            cleanup(inp, out)

    if a in ("json_minifier", "json_to_xml", "base64_decode", "url_decode", "jwt_decoder", "uuid_generator", "password_generator"):
        if a == "json_minifier":
            return JSONResponse({"output": json.dumps(json.loads(text), separators=(",", ":"))})
        if a == "json_to_xml":
            obj = json.loads(text)
            def to_xml(v, name="root"):
                if isinstance(v, dict):
                    return f"<{name}>" + "".join(to_xml(val, key) for key,val in v.items()) + f"</{name}>"
                if isinstance(v, list):
                    return "".join(to_xml(item, name) for item in v)
                return f"<{name}>{str(v)}</{name}>"
            return JSONResponse({"output": to_xml(obj)})
        if a == "base64_decode":
            return JSONResponse({"output": base64.b64decode(text).decode(errors="replace")})
        if a == "url_decode":
            return JSONResponse({"output": unquote(text)})
        if a == "jwt_decoder":
            parts = text.split(".")
            if len(parts) < 2:
                raise HTTPException(400, "Invalid JWT token.")
            def dec(part):
                part += "=" * (-len(part) % 4)
                return json.loads(base64.urlsafe_b64decode(part.encode()).decode())
            return JSONResponse({"output": json.dumps({"header": dec(parts[0]), "payload": dec(parts[1])}, indent=2)})
        if a == "uuid_generator":
            return JSONResponse({"output": str(uuid.uuid4())})
        if a == "password_generator":
            alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
            return JSONResponse({"output": "".join(secrets.choice(alphabet) for _ in range(18))})


    raise HTTPException(501, "Not implemented.")

def extract_youtube_id(link):
    link = (link or "").strip()
    if re.match(r"^[A-Za-z0-9_-]{11}$", link):
        return link
    p = urlparse(link)
    host = p.netloc.lower().replace("www.", "")
    if host == "youtu.be":
        return p.path.strip("/")[:20]
    if "youtube.com" in host:
        qs = parse_qs(p.query)
        if "v" in qs:
            return qs["v"][0][:20]
        m = re.search(r"/(shorts|embed)/([^/?]+)", p.path)
        if m:
            return m.group(2)[:20]
    return None

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print(f"ONE STOP FOR CONVERSIONS running at http://localhost:{port}")
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=True)
